#!/usr/bin/env python3
"""Generate a CV .docx from structured career data.

    python3 engine/lib/build_cv.py -o out.docx
    python3 engine/lib/build_cv.py --career profile/career.json --tailor tailor.json -o out.docx

Why this exists
---------------
The old build copied the previous application's CV and edited it. That
preserved formatting but NOT link correctness: press links decayed into
homepages, canonical sources drifted, and half the set silently went missing.
`check_cv.py` was written to detect that. Generating from data removes the
failure mode instead of auditing for it - hyperlinks are rebuilt from
profile/links.json every time, so there is nothing to inherit and nothing to
drift.

How it works
------------
It **writes into the template** rather than constructing a document. The
formatting that makes these files work - the right tab stop that puts dates at
the margin, the numbering definition behind the bullets, the border under each
heading - is fragile and mostly invisible, and rebuilding it from scratch is
how it breaks. So the template supplies structure; this supplies text.

Paragraphs are cloned from real ones in the template, so every generated
paragraph inherits a known-good style rather than a guessed one.

Tailoring
---------
`--tailor` takes a JSON overlay: replace the summary, reorder or select
competencies, swap individual bullets. It never invents content - anything it
introduces still has to survive `check_cv.py` and the profile guardrails.
"""

import argparse
import copy
import json
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import paths  # noqa: E402
from check_cv import normalize_canonical  # noqa: E402

try:
    from docx import Document
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Inches
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("build_cv.py needs python-docx:  pip install python-docx")


# --------------------------------------------------------------------------
# template surgery
# --------------------------------------------------------------------------

def para_after(ref, template_para):
    """Insert a clone of `template_para` immediately after `ref`; return it."""
    new = copy.deepcopy(template_para._p)
    ref._p.addnext(new)
    from docx.text.paragraph import Paragraph
    return Paragraph(new, ref._parent)


def set_text(p, text):
    """Replace a paragraph's text, keeping the first run's formatting."""
    runs = p.runs
    if not runs:
        p.add_run(text)
        return p
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""
    return p


def clear_hyperlinks(p):
    """Drop any inherited hyperlink, leaving its text as a plain run.

    This is the step that kills the drift bug: nothing is carried forward, so
    every link in the output was placed deliberately by this run.
    """
    for h in list(p._p.findall(qn("w:hyperlink"))):
        for r in h.findall(qn("w:r")):
            h.addprevious(r)
        h.getparent().remove(h)


def set_label_body(p, label, body):
    """Render 'Label: body' keeping the prototype's bold label.

    These lines are two runs - a bold label and a normal body. Collapsing them
    into run 0 makes the WHOLE line bold, and bold is wider, so it wraps
    earlier and silently costs lines. It cost two, and pushed a 1.57-page CV to
    1.61 - over the cap, from a formatting slip invisible in any text diff.
    """
    runs = p.runs
    if label and len(runs) >= 2:
        runs[0].text = f"{label}:"
        runs[1].text = f" {body}"
        for r in runs[2:]:
            r.text = ""
        return p
    return set_text(p, f"{label}: {body}" if label else body)


def _char_formats(p):
    """[(char, rPr)] for every character, so mixed formatting survives a rebuild."""
    out = []
    for r in p.runs:
        rpr = r._element.find(qn("w:rPr"))
        for ch in (r.text or ""):
            out.append((ch, rpr))
    return out


def _base_rpr(p):
    """A copy of the paragraph's first run properties, to inherit face and size."""
    if not p.runs:
        return None
    src = p.runs[0]._element.find(qn("w:rPr"))
    return copy.deepcopy(src) if src is not None else None


def _text_run(rpr, text):
    from docx.oxml import OxmlElement
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    return r


def _link_element(p, rpr, anchor, url):
    from docx.oxml import OxmlElement
    rid = p.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    r = OxmlElement("w:r")
    lp = copy.deepcopy(rpr) if rpr is not None else OxmlElement("w:rPr")
    for tag, val in (("w:color", "0563C1"), ("w:u", "single")):
        for old in lp.findall(qn(tag)):
            lp.remove(old)
        el = OxmlElement(tag)
        el.set(qn("w:val"), val)
        lp.append(el)
    r.append(lp)
    t = OxmlElement("w:t")
    t.text = anchor
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    link.append(r)
    return link


def apply_links(p, canonical, already):
    """Rebuild a paragraph as text/link/text/link… in document order.

    Handles the case that broke the naive version: **two anchors in one
    paragraph.** Appending each link to the end of the paragraph reorders the
    text and corrupts it, which is invisible in the XML and obvious in the
    rendered page. So find every anchor first, then rebuild the paragraph once,
    left to right.
    """
    full = "".join(r.text or "" for r in p.runs)
    if not full:
        return []

    hits = []
    for anchor, url in canonical:
        if anchor in already:
            continue
        i = full.find(anchor)
        if i >= 0:
            hits.append((i, i + len(anchor), anchor, url))
    if not hits:
        return []

    hits.sort()
    kept, end = [], -1
    for h in hits:                       # drop overlaps, e.g. "GitHub" inside a longer anchor
        if h[0] >= end:
            kept.append(h)
            end = h[1]
    if not kept:
        return []

    # Preserve per-character formatting: a paragraph may mix a bold label with
    # normal body text, and flattening it onto one rPr changes the line's width.
    chars = _char_formats(p)
    for r in list(p.runs):
        r._element.getparent().remove(r._element)

    def emit(lo, hi):
        """Emit runs for full[lo:hi], grouping consecutive same-format chars."""
        i = lo
        while i < hi:
            rpr = chars[i][1] if i < len(chars) else None
            j = i
            while j < hi and (chars[j][1] if j < len(chars) else None) is rpr:
                j += 1
            p._p.append(_text_run(copy.deepcopy(rpr) if rpr is not None else None,
                                  full[i:j]))
            i = j

    cursor = 0
    for start, stop, anchor, url in kept:
        if start > cursor:
            emit(cursor, start)
        at = chars[start][1] if start < len(chars) else None
        p._p.append(_link_element(p, copy.deepcopy(at) if at is not None else None,
                                  anchor, url))
        cursor = stop
    if cursor < len(full):
        emit(cursor, len(full))

    return [a for _, _, a, _ in kept]


# --------------------------------------------------------------------------
# build
# --------------------------------------------------------------------------

HYPERLINK_RELTYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)


def prune_orphan_link_rels(doc):
    """Delete hyperlink relationships no longer referenced by any w:hyperlink.

    Removing a link element does not remove its relationship, so the template's
    placeholder URLs survive in word/_rels/document.xml.rels as orphans. They
    are invisible in the document and invisible to a check that walks
    w:hyperlink elements - but the standard link audit greps the rels file, so
    it would report example.com placeholders in a finished CV. Prune them.
    """
    used = {h.get(qn("r:id")) for h in doc.element.body.iter(qn("w:hyperlink"))}
    used.discard(None)
    rels = doc.part.rels
    dropped = [rid for rid, rel in list(rels.items())
               if rel.reltype == HYPERLINK_RELTYPE and rid not in used]
    for rid in dropped:
        del rels[rid]
    return dropped


def find_heading(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    return None


def harvest_protos(doc, heads):
    """Collect one prototype paragraph per role, from the template's own layout.

    Keyed by the role the paragraph plays - summary, competencies, dated (a
    role or institution line), company, bullet, additional - so each generated
    paragraph inherits the spacing of the paragraph it is actually replacing
    rather than of some arbitrary other one.
    """
    SUMMARY, COMPS, EXP, EDU, ADD = heads
    protos, section, seen_role = {}, None, False

    for p in doc.paragraphs:
        text = p.text.strip()
        if text in heads:
            section = text
            continue
        if not text:
            continue

        style = (p.style.name if p.style is not None else "") or ""
        is_bullet = style.replace(" ", "") == "ListParagraph"
        is_dated = "\t" in p.text

        if section == SUMMARY:
            protos.setdefault("summary", p)
        elif section == COMPS:
            protos.setdefault("competencies", p)
        elif section == EXP:
            if is_bullet:
                protos.setdefault("bullet", p)
            elif is_dated:
                protos.setdefault("dated", p)
                seen_role = True
            elif seen_role:
                protos.setdefault("company", p)
        elif section == EDU:
            # BOTH education lines carry a tab, so "is it dated?" cannot tell
            # them apart. Order does: the first is the institution (bold), the
            # second is its degree detail (not bold). Using the institution
            # prototype for both silently bolds every degree line.
            if "edu_institution" not in protos:
                protos["edu_institution"] = p
            elif "edu_detail" not in protos:
                protos["edu_detail"] = p
        elif section == ADD:
            protos.setdefault("additional", p)

    # Sensible fallbacks so an unusual template still builds.
    protos.setdefault("edu_institution", protos.get("dated"))
    protos.setdefault("edu_detail", protos.get("edu_institution"))
    protos.setdefault("additional", protos.get("summary"))
    protos.setdefault("competencies", protos.get("summary"))
    return protos


def build(career, cfg, links, tailor, template_path, out_path):
    doc = Document(template_path)
    heads = cfg["section_order"]

    # A right tab stop at the text width, DERIVED - never hardcoded, because a
    # literal position silently mismatches the moment margins change.
    # (subtracting Length objects yields a plain EMU int, so convert explicitly)
    s = doc.sections[0]
    EMU_PER_INCH = 914400
    tab_in = (s.page_width - s.left_margin - s.right_margin) / EMU_PER_INCH

    # Harvest a prototype paragraph PER ROLE, from the template's own layout.
    #
    # A single generic "plain" prototype is not good enough: the summary, the
    # company line under a role, an education detail, and an ADDITIONAL item
    # all carry different spacing and indentation. Cloning one for all of them
    # inflates the document - it rendered 2.12 pages against a 1.57-page
    # original with byte-identical text. Spacing is invisible in a text diff
    # and only shows up once you measure the PDF.
    protos = harvest_protos(doc, heads)
    missing = [k for k in ("dated", "company", "bullet", "summary") if k not in protos]
    if missing:
        sys.exit(f"template is missing a prototype paragraph for: {', '.join(missing)}")

    def wipe_section(head_text, next_head_text):
        """Delete every paragraph between two headings; return the heading."""
        head = find_heading(doc, head_text)
        if head is None:
            sys.exit(f"template has no {head_text!r} heading")
        nxt = find_heading(doc, next_head_text) if next_head_text else None
        killing = False
        for p in list(doc.paragraphs):
            if p._p is head._p:
                killing = True
                continue
            if nxt is not None and p._p is nxt._p:
                break
            if killing:
                p._p.getparent().remove(p._p)
        return head

    def dated_line(after, left, right, proto="dated"):
        p = para_after(after, protos[proto])
        clear_hyperlinks(p)
        set_text(p, f"{left}\t{right}" if right else left)
        p.paragraph_format.tab_stops.clear_all()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(tab_in), WD_TAB_ALIGNMENT.RIGHT)
        return p

    def plain_line(after, text, proto):
        p = para_after(after, protos[proto])
        clear_hyperlinks(p)
        set_text(p, text)
        p.paragraph_format.tab_stops.clear_all()
        return p

    def bullet_line(after, text):
        p = para_after(after, protos["bullet"])
        clear_hyperlinks(p)
        set_text(p, text)
        return p

    # --- header -----------------------------------------------------------
    hdr = [p for p in doc.paragraphs][:2]
    if hdr:
        clear_hyperlinks(hdr[0])
        set_text(hdr[0], career["header"]["name"])
    if len(hdr) > 1:
        clear_hyperlinks(hdr[1])
        set_text(hdr[1], career["header"]["contact_line"])

    # --- summary ----------------------------------------------------------
    nxt = dict(zip(heads, heads[1:] + [None]))
    head = wipe_section(heads[0], nxt[heads[0]])
    plain_line(head, tailor.get("summary", career["summary"]), "summary")

    # --- competencies -----------------------------------------------------
    comps = tailor.get("competencies", career["competencies"])
    head = wipe_section(heads[1], nxt[heads[1]])
    plain_line(head, "  |  ".join(comps), "competencies")

    # --- experience -------------------------------------------------------
    bullet_over = tailor.get("bullets", {})
    head = wipe_section(heads[2], nxt[heads[2]])
    cur = head
    for role in career["experience"]:            # strictly reverse-chronological
        cur = dated_line(cur, role["title"], role["dates"])
        company = role["company"]
        if role.get("location"):
            company = f"{company}   {role['location']}"
        cur = plain_line(cur, company, "company")
        for i, b in enumerate(role["bullets"]):
            key = f"{role['company']}|{i}"
            cur = bullet_line(cur, bullet_over.get(key, b))

    # --- education --------------------------------------------------------
    head = wipe_section(heads[3], nxt[heads[3]])
    cur = head
    for e in career["education"]:
        cur = dated_line(cur, e["institution"], e.get("date", ""), "edu_institution")
        if e.get("detail"):
            cur = dated_line(cur, e["detail"], e.get("location", ""), "edu_detail")

    # --- additional -------------------------------------------------------
    head = wipe_section(heads[4], nxt[heads[4]])
    cur = head
    for item in career["additional"]:
        cur = para_after(cur, protos["additional"])
        clear_hyperlinks(cur)
        set_label_body(cur, item["label"], item["text"])
        cur.paragraph_format.tab_stops.clear_all()

    # --- hyperlinks: REBUILT from profile/links.json, never inherited ------
    placed = []
    # Longest anchors first: a short anchor nested inside a longer one would
    # otherwise claim the text and strand the longer link.
    canonical = sorted(normalize_canonical(links.get("canonical", {})),
                       key=lambda t: -len(t[0]))
    for p in doc.paragraphs:
        placed.extend(apply_links(p, canonical, placed))

    orphans = prune_orphan_link_rels(doc)
    doc.save(out_path)
    return placed, orphans


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("-o", "--out", required=True)
    ap.add_argument("--career", default=os.path.join(paths.PROFILE_DIR, "career.json"))
    ap.add_argument("--links", default=os.path.join(paths.PROFILE_DIR, "links.json"))
    ap.add_argument("--config", default=os.path.join(paths.CONFIG_DIR, "documents.json"))
    ap.add_argument("--template", default=os.path.join(
        paths.PROJECT_ROOT, "engine", "templates", "cv", "cv_template_01.docx"))
    ap.add_argument("--tailor", help="JSON overlay: summary, competencies, bullets")
    args = ap.parse_args()

    for label, path in (("career", args.career), ("links", args.links),
                        ("config", args.config), ("template", args.template)):
        if not os.path.exists(path):
            sys.exit(f"{label} not found: {path}")

    career = json.load(open(args.career))
    links = json.load(open(args.links))
    cfg = json.load(open(args.config))
    tailor = json.load(open(args.tailor)) if args.tailor else {}

    placed, orphans = build(career, cfg, links, tailor, args.template, args.out)
    roles = len(career["experience"])
    bullets = sum(len(r["bullets"]) for r in career["experience"])
    print(f"wrote {args.out}: {roles} roles, {bullets} bullets, "
          f"{len(placed)} hyperlinks rebuilt from {os.path.basename(args.links)}")
    print(f"  anchors: {', '.join(placed)}")
    if orphans:
        print(f"  pruned {len(orphans)} orphaned template link relationship(s)")
    hi = cfg.get("links", {}).get("max")
    if hi and len(placed) > hi:
        # Not auto-trimmed: which link to drop is an editorial call, and making
        # it silently is exactly the kind of invisible edit this tool avoids.
        print(f"  note: {len(placed)} links exceeds the configured target of {hi} "
              f"- drop the least role-relevant ones by hand")
    print("  now audit it:  python3 engine/lib/check_cv.py --cv " + args.out)


if __name__ == "__main__":
    main()
