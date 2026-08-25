#!/usr/bin/env python3
"""Generate a cover letter .docx from the profile plus a per-application overlay.

    python3 engine/lib/build_cover_letter.py --letter letter.json -o out.docx

The overlay (`--letter`) is the only per-application input:

    {
      "company":   "Acme",
      "role":      "Director of Product",
      "date":      "August 24, 2026",          optional, defaults to today
      "salutation":"Dear Hiring Team,",        optional
      "paragraphs": ["…", "…", "…"]            3-4 short paragraphs
    }

Like the CV generator, it writes INTO the template so the header, font, and
spacing are inherited rather than rebuilt, and it rebuilds hyperlinks from
profile/links.json rather than carrying any forward.

It does not write your paragraphs for you. Deciding what to say about a
specific role is the judgement this tool deliberately leaves to a human or to
the model driving it - what it guarantees is that whatever gets said lands in a
correctly formatted document, on one page, with the guardrails enforced.
"""

import argparse
import datetime
import json
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import paths  # noqa: E402
from check_cv import normalize_canonical  # noqa: E402
from build_cv import (  # noqa: E402
    apply_links, clear_hyperlinks, para_after, prune_orphan_link_rels, set_text,
)

try:
    from docx import Document
except ImportError:
    sys.exit("build_cover_letter.py needs python-docx:  pip install python-docx")


def build(letter, identity, links, template_path, out_path):
    doc = Document(template_path)
    paras = [p for p in doc.paragraphs]
    if len(paras) < 10:
        sys.exit("cover letter template has an unexpected shape")

    name = identity["name"]["full"].upper()
    contact = letter.get("contact_line") or identity.get("contact_line") or ""
    if not contact:
        c = identity.get("contact", {})
        bits = [c.get("location"), c.get("email"), "LinkedIn"]
        contact = "  |  ".join(b for b in bits if b)

    date = letter.get("date") or datetime.date.today().strftime("%B %-d, %Y")
    body = [p for p in letter.get("paragraphs", []) if p.strip()]
    if not body:
        sys.exit("--letter supplied no paragraphs; nothing to write")

    # The template's fixed frame: name, contact, date, recipient, company,
    # salutation, [body…], sign-off, signature.
    fixed = [
        (0, name),
        (1, contact),
        (2, date),
        (3, letter.get("recipient", "Hiring Team")),
        (4, letter["company"]),
        (5, letter.get("salutation", "Dear Hiring Team,")),
    ]
    for idx, text in fixed:
        clear_hyperlinks(paras[idx])
        set_text(paras[idx], text)

    # Body paragraphs 6..8 in the template; clone the last one if more are needed.
    body_slots = paras[6:9]
    proto = body_slots[-1]
    for i, text in enumerate(body[:len(body_slots)]):
        clear_hyperlinks(body_slots[i])
        set_text(body_slots[i], text)
    cur = body_slots[min(len(body), len(body_slots)) - 1]
    for text in body[len(body_slots):]:
        cur = para_after(cur, proto)
        clear_hyperlinks(cur)
        set_text(cur, text)
    # Blank any body slot the letter did not fill, so lorem ipsum cannot ship.
    for leftover in body_slots[len(body):]:
        clear_hyperlinks(leftover)
        set_text(leftover, "")

    tail = [p for p in doc.paragraphs][-2:]
    clear_hyperlinks(tail[0])
    set_text(tail[0], letter.get("signoff", "Best regards,"))
    clear_hyperlinks(tail[1])
    set_text(tail[1], identity["name"]["full"])

    placed = []
    canonical = sorted(normalize_canonical(links.get("canonical", {})),
                       key=lambda t: -len(t[0]))
    for p in doc.paragraphs:
        placed.extend(apply_links(p, canonical, placed))

    orphans = prune_orphan_link_rels(doc)
    doc.save(out_path)
    return placed, orphans, len(body)


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("-o", "--out", required=True)
    ap.add_argument("--letter", required=True, help="per-application JSON overlay")
    ap.add_argument("--identity", default=os.path.join(paths.PROFILE_DIR, "identity.json"))
    ap.add_argument("--links", default=os.path.join(paths.PROFILE_DIR, "links.json"))
    ap.add_argument("--template", default=os.path.join(
        paths.PROJECT_ROOT, "engine", "templates", "cover_letter",
        "cover_letter_template_02.docx"))
    args = ap.parse_args()

    for label, path in (("letter", args.letter), ("identity", args.identity),
                        ("links", args.links), ("template", args.template)):
        if not os.path.exists(path):
            sys.exit(f"{label} not found: {path}")

    letter = json.load(open(args.letter))
    for required in ("company", "paragraphs"):
        if required not in letter:
            sys.exit(f"--letter is missing required key: {required}")

    placed, orphans, n = build(letter, json.load(open(args.identity)),
                               json.load(open(args.links)), args.template, args.out)
    words = sum(len(p.split()) for p in letter["paragraphs"])
    print(f"wrote {args.out}: {n} paragraphs, ~{words} words, {len(placed)} hyperlink(s)")
    if orphans:
        print(f"  pruned {len(orphans)} orphaned template link relationship(s)")
    if not 250 <= words <= 400:
        print(f"  note: {words} words is outside the 250-400 target for a cover letter")
    print("  now audit it:  python3 engine/lib/check_cv.py --cv <the CV> "
          "(letter checks run when both sit in one folder)")


if __name__ == "__main__":
    main()
